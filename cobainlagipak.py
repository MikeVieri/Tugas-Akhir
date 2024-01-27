from docx import Document

def read_txt_file(file_path):
    with open(file_path, 'r') as file:
        content = file.read()
    return content

def write_to_word(content, output_path):
    document = Document()
    document.add_paragraph(content)
    document.save(output_path)

def parse_design_thinking_artifacts(design_thinking_files):
    parsed_data = {}

    for file_path in design_thinking_files:
        content = read_txt_file(file_path)

        # Parse design thinking content based on conditions
        if "POV:" in content:
            parsed_data['Purpose'] = content.split('POV:')[1].strip()
        elif "How might we:" in content:
            parsed_data['Scope'] = content.split('How might we:')[1].strip()
        elif "Site Map:" in content:
            parsed_data['ProductFunction'] = content.split('Site Map:')[1].strip()
        elif "Empathy Map:" in content:
            parsed_data['UserCharacteristic'] = content.split('Empathy Map:')[1].strip()
        elif "User Flow:" in content:
            parsed_data['FunctionalRequirement'] = content.split('User Flow:')[1].strip()

    return parsed_data

def generate_srs_document(parsed_data, output_path):
    document = Document()

    # Add IEEE 830 SRS outline
    document.add_heading('1. Introduction', level=1)
    
    # Add Purpose section
    document.add_heading('1.1 Purpose', level=2)
    document.add_paragraph(parsed_data.get('Purpose', ''))
    
    # Add Scope section
    document.add_heading('1.2 Scope', level=2)
    document.add_paragraph(parsed_data.get('Scope', ''))

    # Add Product Function section
    document.add_heading('2. Overall Description', level=1)
    document.add_heading('2.1 Product Function', level=2)
    document.add_paragraph(parsed_data.get('ProductFunction', ''))

    # Add User Characteristic section
    document.add_heading('3. Specific Requirements', level=1)
    document.add_heading('3.1 User Characteristics', level=2)
    # document.add_paragraph(parsed_data.get('UserCharacteristic', ''))

    #Add Functional Requirement section
    document.add_heading('3.2 Functional Requirements', level=2)
    document.add_paragraph(parsed_data.get('FunctionalRequirement', ''))

    # Save the document
    document.save(output_path)

if __name__ == "__main__":
    # List of paths to design thinking artifact text files
    design_thinking_files = ["pov.txt", "how_might_we.txt", 
                             "site_map.txt", "empathy_map.txt", "user_flow.txt"]

    # Output path for the generated SRS document
    srs_output_path = "SRS_Document.docx"

    # Parse design thinking artifacts
    parsed_data = parse_design_thinking_artifacts(design_thinking_files)

    # Generate SRS document
    generate_srs_document(parsed_data, srs_output_path)

    print(f"SRS document created at {srs_output_path}")
